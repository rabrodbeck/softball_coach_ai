# make_roadmap_doc.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading_with_spacing(doc, text, level, before=12, after=6):
    """Adds a heading with explicit paragraph spacing."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before)
    heading.paragraph_format.space_after = Pt(after)
    heading.paragraph_format.keep_with_next = True
    
    # Color headings to match corporate Navy/Teal style
    run = heading.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Deep Navy
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Teal/Blue
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return heading

def add_code_block(doc, code_text):
    """Creates a beautifully styled code block table in Word."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F2F4F7") # Soft gray
    
    # Left border color simulation via indent
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    
    # Empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_callout_box(doc, title, text, type_box="TIP"):
    """Creates a premium styled callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    
    bg_color = "FFF2CC" if type_box == "WARNING" else "E2EFDA" # Gold vs Pale Green
    border_color = "D9E1F2"
    set_cell_background(cell, bg_color)
    
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run_title = p.add_run(f"★ {title}: ")
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x38, 0x57, 0x23) if type_box == "TIP" else RGBColor(0x7F, 0x60, 0x00)
    
    run_text = p.add_run(text)
    run_text.font.italic = True
    
    # Spacing after
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

# Initialize Document
doc = Document()

# Define Margins (1 inch all around)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure Normal Text Style
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- COVER BANNER ---
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(12)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("Softball Coach AI")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(24)
sub_run = subtitle.add_run("Full-Stack React, TypeScript, FastAPI & Supabase Migration Workbook")
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# --- LEAD DEV BANNER ---
add_callout_box(
    doc, 
    "Lead Developer Note", 
    "This workbook serves as your structured tutorial guide. Work through each milestone step-by-step. Do not rush, and always test each component before moving to the next layer.",
    type_box="TIP"
)

# --- SECTION 1: ARCHITECTURE ---
add_heading_with_spacing(doc, "Part 1: Architectural Evolution", level=1, before=18, after=12)

p_intro = doc.add_paragraph(
    "By moving away from Streamlit, you are transitioning your application from a single-threaded server monolith "
    "to a modern decoupled client-server architecture. This eliminates script rerun lag, ensures strict safety "
    "for API secrets, and introduces a highly scalable API structure."
)
p_intro.paragraph_format.space_after = Pt(8)

# --- SECTION 2: MILESTONE 1 ---
add_heading_with_spacing(doc, "Milestone 1: Supabase Setup & Postgres Database", level=1, before=24, after=12)

doc.add_paragraph(
    "First, you will set up a free cloud-hosted database on Supabase to keep coach profiles and vector chunks "
    "persistent, solving the issue of serverless ephemeral drives wiping local SQLite files."
)

p_steps = doc.add_paragraph()
p_steps.paragraph_format.left_indent = Inches(0.25)
p_steps.paragraph_format.space_after = Pt(6)
p_steps.add_run("1. Create a project on ").font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p_steps.add_run("Supabase.com").bold = True
p_steps.add_run(" (free tier).\n")
p_steps.add_run("2. Enable the pgvector extension in the Supabase SQL editor using:\n")

add_code_block(doc, "CREATE EXTENSION IF NOT EXISTS vector;")

doc.add_paragraph("3. Create the Database Schema tables for both profiles and LangChain embeddings:")

add_code_block(doc, 
    "-- Create coaches profiles\n"
    "CREATE TABLE IF NOT EXISTS coaches (\n"
    "    id SERIAL PRIMARY KEY,\n"
    "    username VARCHAR(255) UNIQUE NOT NULL,\n"
    "    password_hash VARCHAR(255) NOT NULL,\n"
    "    coach_name VARCHAR(255) NOT NULL,\n"
    "    location VARCHAR(255) NOT NULL,\n"
    "    primary_age_group VARCHAR(50) NOT NULL,\n"
    "    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP\n"
    ");\n\n"
    "-- Create RAG embedding database table\n"
    "CREATE TABLE IF NOT EXISTS langchain_pg_embedding (\n"
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
    "    collection_id UUID,\n"
    "    embedding VECTOR(1536),\n"
    "    document VARCHAR,\n"
    "    cmetadata JSONB,\n"
    "    custom_id VARCHAR\n"
    ");"
)

# --- SECTION 3: REFACTOR DATABASE LAYER ---
add_heading_with_spacing(doc, "Phase 2: Adapting your Database Code", level=2, before=18, after=10)
doc.add_paragraph(
    "Modify your local src/database.py file to connect directly to the Supabase Cloud PostgreSQL database using "
    "psycopg2 instead of sqlite3. Ensure you hold credentials in a secure .env file locally:"
)

add_code_block(doc,
    "# src/database.py\n"
    "import psycopg2\n"
    "from psycopg2.extras import RealDictCursor\n"
    "import hashlib\n"
    "import re\n"
    "import os\n"
    "from dotenv import load_dotenv\n\n"
    "load_dotenv()\n\n"
    "DATABASE_URL = os.getenv(\"DATABASE_URL\")\n\n"
    "def get_db_connection():\n"
    "    return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)\n\n"
    "def register_coach(username, password, coach_name, location, age_group):\n"
    "    conn = get_db_connection()\n"
    "    cursor = conn.cursor()\n"
    "    try:\n"
    "        pwd_hash = hash_password(password)\n"
    "        cursor.execute('''\n"
    "            INSERT INTO coaches (username, password_hash, coach_name, location, primary_age_group)\n"
    "            VALUES (%s, %s, %s, %s, %s)\n"
    "            ''', (username.lower().strip(), pwd_hash, coach_name.strip(), location.strip(), age_group))\n"
    "        conn.commit()\n"
    "        return True\n"
    "    except psycopg2.IntegrityError:\n"
    "        return False\n"
    "    finally:\n"
    "        cursor.close()\n"
    "        conn.close()"
)

# --- SECTION 4: FASTAPI WEB API ---
add_heading_with_spacing(doc, "Milestone 2: Constructing the FastAPI Web API Core", level=1, before=24, after=12)
doc.add_paragraph(
    "FastAPI maps backend Python modules to clear HTTP interfaces that can be called asynchronously by React. "
    "Install the requirements (fastapi, uvicorn) and construct your server wrapper:"
)

add_code_block(doc,
    "# src/main.py\n"
    "from fastapi import FastAPI, HTTPException\n"
    "from fastapi.middleware.cors import CORSMiddleware\n"
    "from pydantic import BaseModel, EmailStr\n"
    "from src.database import authenticate_coach, register_coach\n\n"
    "app = FastAPI(title=\"Softball Coach AI API\")\n\n"
    "app.add_middleware(\n"
    "    CORSMiddleware,\n"
    "    allow_origins=[\"*\"], # Swap with your React URL in prod\n"
    "    allow_credentials=True,\n"
    "    allow_methods=[\"*\"],\n"
    "    allow_headers=[\"*\"],\n"
    ")\n\n"
    "class RegisterRequest(BaseModel):\n"
    "    username: EmailStr\n"
    "    password: str\n"
    "    coach_name: str\n"
    "    location: str\n"
    "    age_group: str\n\n"
    "@app.post(\"/api/auth/register\")\n"
    "def api_register(data: RegisterRequest):\n"
    "    success = register_coach(data.username, data.password, data.coach_name, data.location, data.age_group)\n"
    "    if not success:\n"
    "        raise HTTPException(status_code=400, detail=\"Username already exists.\")\n"
    "    return {\"message\": \"Account created successfully!\"}"
)

# --- SECTION 5: FRONTEND REACT ---
add_heading_with_spacing(doc, "Milestone 3: Scaffolding the React & TypeScript Frontend", level=1, before=24, after=12)
doc.add_paragraph(
    "Use Vite to scaffold the frontend project inside a subdirectory. Choose React and TypeScript when prompted:"
)

add_code_block(doc,
    "npm create vite@latest frontend -- --template react-ts\n"
    "cd frontend\n"
    "npm install\n"
    "npm install lucide-react"
)

doc.add_paragraph(
    "This creates a clean, compiled setup. You will build highly responsive UI components and fetch API endpoints "
    "using async/await methods, replacing the rigid top-to-bottom run mechanics of Streamlit."
)

# --- LEAD DEV TIP ---
add_callout_box(
    doc,
    "Developer Safety Rule",
    "Ensure you never hardcode API keys or database URLs directly in your code. Always reference process.env or os.environ, and check that your local .env is ignored in your Git settings.",
    type_box="WARNING"
)

# Save Document
output_path = "Softball_Coach_AI_React_Migration_Roadmap.docx"
doc.save(output_path)
print(f"Success! Elegant roadmap document written to: {os.path.abspath(output_path)}")
