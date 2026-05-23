# make_doc.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading_with_spacing(doc, text, level, before=12, after=6):
    """Adds a heading with explicit paragraph spacing."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before)
    heading.paragraph_format.space_after = Pt(after)
    heading.paragraph_format.keep_with_next = True
    return heading

# Initialize Document
doc = Document()

# Define Global Styles & Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure Normal Text Style
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- DOCUMENT TITLE ---
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("Integrated Multi-Tenant & Guest Upgrade Specification")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Deep Navy

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(24)
sub_run = subtitle.add_run("Project: Softball Coach AI Command Center")
sub_run.font.size = Pt(14)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# --- CALLOUT BOX: PROJECT GOAL ---
table_goal = doc.add_table(rows=1, cols=1)
table_goal.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_goal = table_goal.cell(0, 0)
set_cell_background(cell_goal, "F2F4F7") # Light Grey/Blue Tint
p_goal = cell_goal.paragraphs[0]
p_goal.paragraph_format.left_indent = Inches(0.15)
p_goal.paragraph_format.right_indent = Inches(0.15)
p_goal.paragraph_format.space_before = Pt(8)
p_goal.paragraph_format.space_after = Pt(8)
run_goal_lbl = p_goal.add_run("Project Goal: ")
run_goal_lbl.bold = True
p_goal.add_run("Establish a local SQLite storage architecture, configure a clear three-way entry portal (Log In, Register, or Continue as Guest), and implement dynamic, context-aware prompt processing for your fastpitch assistant application.")

# --- SECTION 1 ---
add_heading_with_spacing(doc, "Part 1: File Architecture Overview", level=1, before=24, after=12)
p_arch = doc.add_paragraph("Ensure your local directory matches this structure before modifying files:")
p_arch.paragraph_format.space_after = Pt(6)

# Code Block styling for folder structure
table_struct = doc.add_table(rows=1, cols=1)
cell_struct = table_struct.cell(0, 0)
set_cell_background(cell_struct, "F8F9FA")
p_struct = cell_struct.paragraphs[0]
p_struct.paragraph_format.space_before = Pt(6)
p_struct.paragraph_format.space_after = Pt(6)
run_struct = p_struct.add_run(
    "softball_coach_ai/\n"
    "├── .streamlit/\n"
    "│   └── config.toml\n"
    "└── src/\n"
    "    ├── app.py           <-- Replace entirely\n"
    "    ├── database.py      <-- Create brand new\n"
    "    └── retriever.py"
)
run_struct.font.name = 'Consolas'
run_struct.font.size = Pt(9.5)

# --- SECTION 2 ---
add_heading_with_spacing(doc, "Part 2: Step-by-Step Code Implementation", level=1, before=24, after=12)

add_heading_with_spacing(doc, "Phase 1: Creating the Database Layer", level=2, before=14, after=6)
doc.add_paragraph("Create a new file named database.py inside your src/ directory and add the following code to manage authentication and user state mapping:")

# DB Code Table
table_db = doc.add_table(rows=1, cols=1)
cell_db = table_db.cell(0, 0)
set_cell_background(cell_db, "F8F9FA")
p_db = cell_db.paragraphs[0]
p_db.paragraph_format.space_before = Pt(6)
p_db.paragraph_format.space_after = Pt(6)
run_db_code = p_db.add_run(
    "# src/database.py\n"
    "import sqlite3\n"
    "import hashlib\n\n"
    "def get_db_connection():\n"
    "    \"\"\"Establishes connection to the local SQLite database file.\"\"\"\n"
    "    conn = sqlite3.connect(\"softball_app.db\")\n"
    "    conn.row_factory = sqlite3.Row\n"
    "    return conn\n\n"
    "def init_db():\n"
    "    \"\"\"Creates the coaches table automatically if it does not exist.\"\"\"\n"
    "    conn = get_db_connection()\n"
    "    cursor = conn.cursor()\n"
    "    cursor.execute('''\n"
    "        CREATE TABLE IF NOT EXISTS coaches (\n"
    "            id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "            username TEXT UNIQUE NOT NULL,\n"
    "            password_hash TEXT NOT NULL,\n"
    "            coach_name TEXT NOT NULL,\n"
    "            location TEXT NOT NULL,\n"
    "            primary_age_group TEXT NOT NULL\n"
    "        )\n"
    "    ''')\n"
    "    conn.commit()\n"
    "    conn.close()\n\n"
    "def hash_password(password):\n"
    "    \"\"\"Converts plain-text passwords into a secure SHA-256 hash string.\"\"\"\n"
    "    return hashlib.sha256(password.encode()).hexdigest()\n\n"
    "def register_coach(username, password, coach_name, location, age_group):\n"
    "    \"\"\"Attempts to insert a new coach profile into the SQLite database.\"\"\"\n"
    "    conn = get_db_connection()\n"
    "    cursor = conn.cursor()\n"
    "    try:\n"
    "        pwd_hash = hash_password(password)\n"
    "        cursor.execute('''\n"
    "            INSERT INTO coaches (username, password_hash, coach_name, location, primary_age_group)\n"
    "            VALUES (?, ?, ?, ?, ?)\n"
    "        ''', (username.lower().strip(), pwd_hash, coach_name.strip(), location.strip(), age_group))\n"
    "        conn.commit()\n"
    "        return True\n"
    "    except sqlite3.IntegrityError:\n"
    "        return False\n"
    "    finally:\n"
    "        conn.close()\n\n"
    "def authenticate_coach(username, password):\n"
    "    \"\"\"Validates credentials against hashed database entries.\"\"\"\n"
    "    conn = get_db_connection()\n"
    "    cursor = conn.cursor()\n"
    "    pwd_hash = hash_password(password)\n"
    "    cursor.execute('''\n"
    "        SELECT username, coach_name, location, primary_age_group \n"
    "        FROM coaches \n"
    "        WHERE username = ? AND password_hash = ?\n"
    "    ''', (username.lower().strip(), pwd_hash))\n"
    "    row = cursor.fetchone()\n"
    "    conn.close()\n"
    "    if row:\n"
    "        return {\n"
    "            \"username\": row[\"username\"],\n"
    "            \"coach_name\": row[\"coach_name\"],\n"
    "            \"location\": row[\"location\"],\n"
    "            \"age_group\": row[\"primary_age_group\"]\n"
    "        }\n"
    "    return None"
)
run_db_code.font.name = 'Consolas'
run_db_code.font.size = Pt(9)

# --- Phase 2 ---
add_heading_with_spacing(doc, "Phase 2: Updating the Main Streamlit App", level=2, before=18, after=6)
doc.add_paragraph("Completely replace the contents of src/app.py with the implementation below. This code constructs the 3-button landing routing logic and embeds the session controllers:")

# App Code Table
table_app = doc.add_table(rows=1, cols=1)
cell_app = table_app.cell(0, 0)
set_cell_background(cell_app, "F8F9FA")
p_app = cell_app.paragraphs[0]
p_app.paragraph_format.space_before = Pt(6)
p_app.paragraph_format.space_after = Pt(6)
run_app_code = p_app.add_run(
    "# src/app.py\n"
    "import streamlit as st\n"
    "from retriever import build_chain\n"
    "from database import init_db, authenticate_coach, register_coach\n\n"
    "# --- 1. PAGE CONFIGURATION ---\n"
    "st.set_page_config(page_title=\"Softball Coach AI\", page_icon=\"\U0001F94E\", layout=\"centered\")\n"
    "st.title(\"\U0001F94E Softball Coach AI\")\n"
    "st.caption(\"Your AI-powered fastpitch coaching assistant.\")\n\n"
    "# --- 2. SESSION STATE INITIALIZATION ---\n"
    "if \"access_granted\" not in st.session_state:\n"
    "    st.session_state.access_granted = False\n"
    "if \"active_user\" not in st.session_state:\n"
    "    st.session_state.active_user = None\n"
    "if \"auth_mode\" not in st.session_state:\n"
    "    st.session_state.auth_mode = \"menu\"\n"
    "if \"chain\" not in st.session_state:\n"
    "    with st.spinner(\"Loading knowledge base...\"):\n"
    "        st.session_state.chain = build_chain()\n"
    "if \"messages\" not in st.session_state:\n"
    "    st.session_state.messages = []\n"
    "if \"pending_question\" not in st.session_state:\n"
    "    st.session_state.pending_question = None\n"
    "if \"is_sidebar_action\" not in st.session_state:\n"
    "    st.session_state.is_sidebar_action = False\n\n"
    "init_db()\n\n"
    "# --- 3. THE 3-WAY GATEWAY PORTAL INTERFACE ---\n"
    "if not st.session_state.access_granted:\n"
    "    st.subheader(\"\u26BE Coaching Command Center Portal\")\n"
    "    st.caption(\"Access your personalized dugout files or explore as a guest.\")\n\n"
    "    if st.session_state.auth_mode == \"menu\":\n"
    "        col1, col2, col3 = st.columns(3)\n"
    "        with col1:\n"
    "            if st.button(\"\U0001F511 Log In\", use_container_width=True):\n"
    "                st.session_state.auth_mode = \"login\"\n"
    "                st.rerun()\n"
    "        with col2:\n"
    "            if st.button(\"\U0001F4CB Create Account\", use_container_width=True):\n"
    "                st.session_state.auth_mode = \"register\"\n"
    "                st.rerun()\n"
    "        with col3:\n"
    "            if st.button(\"\U0001F94E Continue as Guest\", use_container_width=True):\n"
    "                st.session_state.access_granted = True\n"
    "                st.session_state.active_user = None\n"
    "                st.rerun()\n\n"
    "    elif st.session_state.auth_mode == \"login\":\n"
    "        st.markdown(\"### \U0001F511 Coach Login\")\n"
    "        login_user = st.text_input(\"Username / Email\", key=\"login_user_input\")\n"
    "        login_pwd = st.text_input(\"Password\", type=\"password\", key=\"login_pwd_input\")\n"
    "        c1, c2 = st.columns(2)\n"
    "        with c1:\n"
    "            if st.button(\"Access Boardroom\", use_container_width=True, type=\"primary\"):\n"
    "                user_record = authenticate_coach(login_user, login_pwd)\n"
    "                if user_record:\n"
    "                    st.session_state.access_granted = True\n"
    "                    st.session_state.active_user = user_record\n"
    "                    st.rerun()\n"
    "                else:\n"
    "                    st.error(\"Invalid credentials.\")\n"
    "        with c2:\n"
    "            if st.button(\"\u2B05\ufe0f Back to Portal\", use_container_width=True):\n"
    "                st.session_state.auth_mode = \"menu\"\n"
    "                st.rerun()\n\n"
    "    elif st.session_state.auth_mode == \"register\":\n"
    "        st.markdown(\"### \U0001F4CB Create Your Coaching Profile\")\n"
    "        new_user = st.text_input(\"Choose Username / Email\", key=\"reg_user_input\")\n"
    "        new_pwd = st.text_input(\"Choose Password\", type=\"password\", key=\"reg_pwd_input\")\n"
    "        new_name = st.text_input(\"Coach Full Name\", placeholder=\"Coach Ryan\")\n"
    "        new_loc = st.text_input(\"Your Location\", placeholder=\"Streamwood, IL\")\n"
    "        new_age = st.selectbox(\"Primary Age Group Coached\", [\"8U Division\", \"10U Division\", \"12U Division\", \"14U Division\"])\n"
    "        c1, c2 = st.columns(2)\n"
    "        with c1:\n"
    "            if st.button(\"Build Playbook Account\", use_container_width=True, type=\"primary\"):\n"
    "                if new_user and new_pwd and new_name and new_loc:\n"
    "                    if register_coach(new_user, new_pwd, new_name, new_loc, new_age):\n"
    "                        st.session_state.auth_mode = \"login\"\n"
    "                        st.rerun()\n"
    "                    else:\n"
    "                        st.error(\"Username already taken.\")\n"
    "        with c2:\n"
    "            if st.button(\"\u2B05\ufe0f Back to Portal\", use_container_width=True):\n"
    "                st.session_state.auth_mode = \"menu\"\n"
    "                st.rerun()\n\n"
    "else:\n"
    "    # --- 4. SIDEBAR CONFIGURATION ---\n"
    "    with st.sidebar:\n"
    "        st.header(\"\U0001F4CB Practice Plans\")\n"
    "        age_options = [\"8U Division\", \"10U Division\", \"12U Division\", \"14U Division\"]\n"
    "        if st.session_state.active_user is not None:\n"
    "            user_default_age = st.session_state.active_user[\"age_group\"]\n"
    "            default_index = age_options.index(user_default_age) if user_default_age in age_options else 2\n"
    "        else:\n"
    "            default_index = 2\n"
    "        selected_division = st.selectbox(\"Select Division:\", age_options, index=default_index)\n"
    "        if st.button(\"\U0001F3CF Generate Playbook\", use_container_width=True):\n"
    "            st.session_state.pending_question = f\"Build a comprehensive practice plan template for a {selected_division} team.\"\n"
    "            st.session_state.is_sidebar_action = True\n\n"
    "    # --- 5. RENDER CONVERSATION HISTORY ---\n"
    "    if not st.session_state.messages:\n"
    "        st.info(\"\U0001F4CB Coach's Whiteboard Active: Ask a coaching question or choose a template on the left.\")\n"
    "    for msg in st.session_state.messages:\n"
    "        with st.chat_message(msg[\"role\"], avatar=\"\U0001F4E2\" if msg[\"role\"] == \"user\" else \"\U0001F94E\"):\n"
    "            st.markdown(msg[\"content\"])\n\n"
    "    # --- 6. CAPTURE & PROCESS USER INPUT ---\n"
    "    chat_input_val = st.chat_input(\"Ask a coaching question...\")\n"
    "    prompt = None\n"
    "    is_sidebar = False\n"
    "    if st.session_state.pending_question:\n"
    "        prompt = st.session_state.pending_question\n"
    "        is_sidebar = st.session_state.is_sidebar_action\n"
    "        st.session_state.pending_question = None\n"
    "        st.session_state.is_sidebar_action = False\n"
    "    elif chat_input_val:\n"
    "        prompt = chat_input_val\n\n"
    "    if prompt:\n"
    "        if not is_sidebar:\n"
    "            st.session_state.messages.append({\"role\": \"user\", \"content\": prompt})\n"
    "            with st.chat_message(\"user\", avatar=\"\U0001F4E2\"):\n"
    "                st.markdown(prompt)\n"
    "        if st.session_state.active_user is not None:\n"
    "            profile = st.session_state.active_user\n"
    "            profile_context_string = f\"Directly advising Coach {profile['coach_name']} based in {profile['location']}. Tailor advice for competitive {profile['age_group']} fastpitch players.\"\n"
    "        else:\n"
    "            profile_context_string = \"Directly advising a youth fastpitch coach. Provide structurally sound coaching advice tailored to youth development frameworks.\"\n"
    "        with st.chat_message(\"assistant\", avatar=\"\U0001F94E\"):\n"
    "            with st.spinner(\"Thinking...\"):\n"
    "                result = st.session_state.chain.invoke({\"profile_context\": profile_context_string, \"question\": prompt})\n"
    "                answer = result[\"answer\"]\n"
    "            st.markdown(answer)\n"
    "        st.session_state.messages.append({\"role\": \"assistant\", \"content\": answer})"
)
run_app_code.font.name = 'Consolas'
run_app_code.font.size = Pt(8.5)

# Save Document
output_filename = "Softball_Coach_AI_Specification.docx"
doc.save(output_filename)
print(f"Success! Perfect Word document created as: {os.path.abspath(output_filename)}")